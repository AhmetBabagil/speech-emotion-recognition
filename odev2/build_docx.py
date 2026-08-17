"""Markdown raporu Word (.docx) belgesine çeviren küçük dönüştürücü (Ödev 2).

    python odev2/build_docx.py
    python odev2/build_docx.py --md odev2/PROJE_ILERLEME_RAPORU_2.md --docx cikti.docx

Neden var? Ödevin teslimi Word/Doc formatında istendiği için, markdown olarak
tutulan rapor (başlıklar, tablolar, görseller, listeler, kod blokları) burada
`python-docx` kütüphanesiyle gerçek bir .docx belgesine dönüştürülür. Markdown
kaynak dosyada kalır; belge her değişiklikten sonra bu betikle yeniden üretilir.

Çalışma mantığı: markdown satır satır okunur ve her satırın "türü" (başlık mı,
tablo mu, liste mi, resim mi, düz metin mi) sırayla kontrol edilir; türe uygun
python-docx öğesi belgeye eklenir. Tam bir markdown ayrıştırıcı değildir —
raporda kullanılan alt küme için yeterlidir.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clean_inline(text: str) -> str:
    """Satır içi markdown işaretlerini temizleyip Word'e uygun düz metin döndürür.

    Yapılan dönüşümler sırasıyla:
      * görsel sözdizimi ``![alt](yol)``  → yalnızca yol metni,
      * bağlantı ``[metin](url)``        → ``metin (url)``,
      * kalın ``**`` / ``__`` işaretleri → silinir,
      * kod işareti `` ` ``              → silinir.
    Word'de biçimlendirme stillerle yapıldığı için işaretlerin kendisi gereksizdir.
    """
    text = re.sub(r"!\[[^\]]*\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    return text.strip()


def split_table_row(line: str) -> list[str]:
    """Bir markdown tablo satırını (`| a | b |`) hücre listesine ayırır.

    Baştaki/sondaki `|` karakterleri atılır, kalan metin `|` ile bölünür ve her
    hücre `clean_inline` ile markdown işaretlerinden temizlenir.
    """
    line = line.strip().strip("|")
    return [clean_inline(cell.strip()) for cell in line.split("|")]


def is_table_sep(line: str) -> bool:
    """Satırın markdown tablo ayraç satırı (`|---|:---:|...`) olup olmadığını söyler.

    Tablonun ikinci satırı böyle bir ayraçtır; bir satırın gerçekten tablo
    başlangıcı olduğuna, ancak ALTINDAKİ satır bu kalıba uyuyorsa karar verilir.
    Düzenli ifade `:?-{3,}:?`: isteğe bağlı iki nokta + en az üç tire (hizalama
    işaretli `:---:` biçimleri de kabul edilir).
    """
    parts = [p.strip() for p in line.strip().strip("|").split("|")]
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts)


def set_cell_font(cell, size=7.2):
    """Bir tablo hücresindeki tüm yazıların punto boyutunu ayarlar.

    Rapordaki tablolar geniş olduğu için küçük punto (7.2) kullanılır; böylece
    tablolar sayfaya taşmadan sığar. python-docx'te yazı biçimi run nesnelerinde
    tutulduğundan hücredeki her paragraf ve run tek tek gezilir.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


def add_table(doc: Document, rows: list[str]) -> None:
    """Markdown tablo satırlarından Word tablosu kurup belgeye ekler.

    ``rows[0]`` başlık satırı, ``rows[1]`` ayraç (atlanır), ``rows[2:]`` veridir.
    Başlık hücreleri kalın yazılır; tüm hücreler küçük puntoya çekilir. Veri
    satırı başlıktan kısa gelirse eksik hücreler boş bırakılır (dizin hatası
    yerine boş hücre tercih edilir). Sonda eklenen boş paragraf, tablonun
    ardından gelen metinle yapışmasını önler.
    """
    header = split_table_row(rows[0])
    body = [split_table_row(r) for r in rows[2:]]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True
    for i, val in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(7.2)
    for row in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            cells[i].text = row[i] if i < len(row) else ""
            set_cell_font(cells[i])
    doc.add_paragraph()


def add_image(doc: Document, md_dir: Path, line: str) -> None:
    """Markdown görsel satırındaki resmi bulup ortalanmış olarak belgeye ekler.

    Görsel yolu markdown dosyasının konumuna (md_dir) göre çözülür, çünkü
    rapordaki yollar göreli yazılır. Dosya diskte yoksa belge üretimi durmaz;
    yerine uyarı metni içeren bir paragraf konur. Genişlik 4.6 inç'e sabitlenir
    ki tüm görseller sayfada tutarlı boyutta görünsün.
    """
    match = re.search(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
    if not match:
        return
    img_path = (md_dir / match.group(1)).resolve()
    if not img_path.exists():
        doc.add_paragraph(f"[Gorsel bulunamadi: {img_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(4.6))


def add_paragraph(doc: Document, text: str) -> None:
    """Düz metni, işaretlerden temizleyip standart aralıklı paragraf olarak ekler.

    Temizlik sonrası boş kalan metin için paragraf açılmaz; 6 punto alt boşluk
    paragraflar arasında okunaklı bir aralık bırakır.
    """
    text = clean_inline(text)
    if text:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)


def build_docx(md_path: Path, docx_path: Path) -> None:
    """Markdown dosyasını baştan sona okuyup eşdeğer .docx belgesini üretir.

    Üç bölümden oluşur:
      1. Sayfa düzeni: dar kenar boşlukları (tablolara yer açmak için),
      2. Stiller: gövde ve başlık yazı tipi/puntoları tek yerden ayarlanır,
      3. Satır satır dönüştürme döngüsü: her satırın türü sırayla denenir
         (kod bloğu → boş satır → başlık → görsel → tablo → madde işareti →
         numaralı liste → düz metin) ve uygun öğe belgeye eklenir.

    Art arda gelen düz metin satırları `para_lines` içinde biriktirilir ve tür
    değişince `flush_para` ile TEK paragraf olarak yazılır — markdown'da tek
    paragrafın birden çok satıra bölünebilmesi bunun sebebidir.
    """
    doc = Document()
    # ---- 1) Sayfa düzeni: dar kenar boşlukları -----------------------------
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    # ---- 2) Stiller: tüm belge Calibri, başlıklar kademeli puntolarla ------
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)

    # ---- 3) Satır satır dönüştürme -----------------------------------------
    lines = md_path.read_text(encoding="utf-8").splitlines()
    md_dir = md_path.parent
    i = 0
    in_code = False          # şu an ``` ile açılmış kod bloğunun içinde miyiz?
    code_lines: list[str] = []   # kod bloğu satırları (kapanınca tek parça yazılır)
    para_lines: list[str] = []   # biriken düz metin satırları (tek paragraf olur)

    def flush_para():
        """Biriken düz metin satırlarını tek paragraf olarak belgeye boşaltır."""
        nonlocal para_lines
        if para_lines:
            add_paragraph(doc, " ".join(para_lines))
            para_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Kod bloğu sınırı: ``` hem açar hem kapatır. Kapanışta biriken satırlar
        # eş aralıklı (Consolas) küçük puntolu tek paragraf olarak yazılır.
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                flush_para()
                in_code = True
            i += 1
            continue

        # Kod bloğu içindeki satırlar olduğu gibi (markdown yorumlanmadan) saklanır.
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Boş satır: paragraf sınırıdır — birikeni yaz, yeni paragrafa hazırlan.
        if not stripped:
            flush_para()
            i += 1
            continue

        # Başlık: baştaki '#' sayısı seviyeyi verir (Word'de en fazla 3 kullanılır).
        if stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            title = clean_inline(stripped[level:].strip())
            doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        # Görsel satırı: ![...](yol)
        if stripped.startswith("![](") or stripped.startswith("!["):
            flush_para()
            add_image(doc, md_dir, stripped)
            i += 1
            continue

        # Tablo: '|' ile başlıyor VE bir sonraki satır ayraç (---) satırıysa.
        # Tüm ardışık '|' satırları toplanıp tek seferde tabloya çevrilir.
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            flush_para()
            table_rows = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_table(doc, table_rows)
            continue

        # Madde işaretli liste: "- " ile başlayan satırlar.
        if stripped.startswith("- "):
            flush_para()
            p = doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Numaralı liste: "1. ", "2. " gibi satırlar; numara Word tarafından
        # yeniden üretileceği için metinden çıkarılır.
        if re.match(r"\d+\.\s+", stripped):
            flush_para()
            item = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(clean_inline(item), style="List Number")
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Hiçbir özel türe uymadı: düz metin olarak paragrafa biriktir.
        para_lines.append(stripped)
        i += 1

    flush_para()  # dosya paragraf ortasında bitmiş olabilir; son birikeni yaz
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    """Komut satırı argümanlarını okuyup markdown → docx dönüşümünü çalıştırır.

    ``--md`` kaynak markdown dosyası, ``--docx`` üretilecek Word dosyasıdır;
    ikisinin de varsayılanı Ödev 2 ilerleme raporudur.
    """
    parser = argparse.ArgumentParser()
    parser.add_argument("--md", default="odev2/PROJE_ILERLEME_RAPORU_2.md")
    parser.add_argument("--docx", default="odev2/PROJE_ILERLEME_RAPORU_2.docx")
    args = parser.parse_args()
    build_docx(Path(args.md), Path(args.docx))
    print(f"DOCX written: {args.docx}")


if __name__ == "__main__":
    main()
