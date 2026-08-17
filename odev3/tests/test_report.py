'''Deney çıktılarından üretilen HTML ve Word raporlarının testleri.

İki seviyede test var:
1. **Birim**: yardımcı satır üreticileri (ablasyon/kalibrasyon/kararlılık/
   etki tabloları) gerçek deney dosyalarından beklenen değerleri okuyor mu;
   HTML/DOCX çiziciler Türkçe karakterleri koruyup içeriği doğru kaçırıyor
   (escape) mu?
2. **Uçtan uca**: build_report, depodaki GERÇEK deney çıktılarından tam
   raporu üretir; testte raporun kritik cümleleri, sayıları, tablo ve görsel
   SAYILARI sabitlenir. Böylece rapor üretimi veya deney çıktıları bozulursa
   test anında haber verir.
'''

import json
from pathlib import Path

import pytest

from odev3.build_report import (
    Heading,
    LinkList,
    Paragraph,
    TableBlock,
    _ablation_rows,
    _calibration_rows,
    _feature_stability_rows,
    _hyperparameter_effect_rows,
    _render_docx,
    _render_html,
    build_report,
)


def test_report_helpers_load_expanded_ablation_and_calibration() -> None:
    '''Satır üreticileri, depodaki gerçek deney dosyalarından beklenen değerleri çıkarmalı.

    Buradaki sabitler (16 ablasyon satırı, T=2.0096/3.2744 sıcaklıkları,
    ECE geçişleri, kararlılık ortalama±std metinleri...) gerçekten koşulmuş
    deneylerin sonuçlarıdır. Test, rapor yardımcıları ile deney çıktıları
    arasındaki "sözleşmeyi" sabitler: format ya da veri değişirse fark edilir.
    '''

    ablation_rows = _ablation_rows(Path('odev3/feature_ablation'))
    summary = json.loads(
        Path('odev3/outputs/summary.json').read_text(encoding='utf-8')
    )
    calibration_rows = _calibration_rows(summary['per_dataset'])
    stability_rows = _feature_stability_rows(Path('odev3/feature_stability'))
    effect_rows = _hyperparameter_effect_rows(Path('odev3/hyperparameter_effects'))

    assert len(ablation_rows) == 16
    assert {row[3] for row in ablation_rows} == {64, 80, 96}
    # Her korpusta tam bir kazanan işaretlenmiş olmalı (toplam 2).
    assert sum(bool(row[-1]) for row in ablation_rows) == 2
    assert len(calibration_rows) == 2
    assert calibration_rows[0][1] == pytest.approx(2.0095735401)
    assert calibration_rows[1][1] == pytest.approx(3.2743607877)
    assert calibration_rows[0][5] == '0.2455 → 0.0575'
    assert calibration_rows[1][5] == '0.1049 → 0.0132'
    assert all(row[-1] is True for row in calibration_rows)
    assert len(stability_rows) == 4
    assert sum(bool(row[-1]) for row in stability_rows) == 2
    assert stability_rows[0][8] == '0.4644 ± 0.0063'
    assert stability_rows[2][8] == '0.2370 ± 0.0117'
    assert len(effect_rows) == 16
    assert effect_rows[0][0:2] == ('CREMA-D', 'Learning rate')
    assert effect_rows[0][4] == 1
    assert effect_rows[0][5] == pytest.approx(0.4684975243)
    assert effect_rows[8][0:2] == ('MELD', 'Learning rate')
    assert effect_rows[8][5] == pytest.approx(0.2414046227)


def test_html_renderer_preserves_turkish_text_and_escapes_content(
    tmp_path: Path,
) -> None:
    '''HTML çizici Türkçe karakterleri korumalı ve tehlikeli karakterleri kaçırmalı.

    Kontroller: UTF-8 meta etiketi, 'Ö/ç/ş' gibi karakterlerin aynen
    kalması, '<' işaretinin '&lt;' olması (HTML enjeksiyonuna karşı), float
    hücrenin 4 basamağa yuvarlanması, tablo altyazısı ve URL'deki '&'nin
    href içinde '&amp;' olarak kaçırılması.
    '''

    output_path = tmp_path / 'report.html'
    blocks = [
        Heading(1, 'Ödev 3 – Türkçe Rapor'),
        Paragraph('MLP giriş boyutu 64×64 ve değer < 5000.'),
        LinkList((('Proje', 'https://example.com/?a=1&b=2'),)),
        TableBlock(
            ('Deney', 'Macro-F1'),
            ((1, 0.45678),),
            'Validation sonuçları',
        ),
    ]

    _render_html(blocks, output_path, 'Ödev 3')
    html = output_path.read_text(encoding='utf-8')

    assert '<meta charset="utf-8">' in html
    assert 'Ödev 3 – Türkçe Rapor' in html
    assert 'değer &lt; 5000' in html
    assert '0.4568' in html
    assert '<caption>Validation sonuçları</caption>' in html
    assert 'href="https://example.com/?a=1&amp;b=2"' in html


def test_docx_renderer_writes_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    '''DOCX çizici başlık, paragraf ve tabloyu Word belgesine doğru aktarmalı.

    Üretilen dosya python-docx ile geri okunur: başlık ve paragraf metinleri
    belgede bulunmalı, tek tablo olmalı ve float hücre '0.5500' biçiminde
    (4 ondalık) yazılmış olmalı.
    '''

    from docx import Document

    output_path = tmp_path / 'report.docx'
    blocks = [
        Heading(1, 'YAP 470 / BİL 570 – Proje İlerleme Raporu 3'),
        Paragraph('Konuşmadan duygu tanıma çalışması.'),
        TableBlock(
            ('Deney', 'Doğruluk', 'Macro-F1'),
            ((1, 0.61, 0.55),),
            'CREMA-D validation sonuçları',
        ),
    ]

    _render_docx(blocks, output_path, 'Proje İlerleme Raporu 3')
    document = Document(output_path)
    paragraph_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)

    assert output_path.stat().st_size > 0
    assert 'Proje İlerleme Raporu 3' in paragraph_text
    assert 'Konuşmadan duygu tanıma çalışması.' in paragraph_text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 2).text == '0.5500'


def test_final_report_builds_from_complete_experiment_artifacts(
    tmp_path: Path,
) -> None:
    '''Uçtan uca: gerçek deney çıktılarından NİHAİ raporun tamamı doğru kurulmalı.

    Bu, rapor üretiminin regresyon kalesidir. Raporun kritik cümleleri
    (koşu sayıları, bootstrap açıklaması, kalibrasyon bulguları, metodolojik
    uyarılar) ve yapısal sayıları (tam 21 tablo, 8 görsel, hiç eksik görsel
    kutusu olmaması) HTML ve DOCX çıktılarının İKİSİNDE birden doğrulanır.
    Deney dosyaları ya da rapor kodu uyumsuz değişirse bu test kırılır.
    '''

    from docx import Document

    html_path = tmp_path / 'final.html'
    docx_path = tmp_path / 'final.docx'

    paths = build_report(
        output_root='odev3/outputs',
        ablation_root='odev3/feature_ablation',
        effect_root='odev3/hyperparameter_effects',
        html_path=html_path,
        docx_path=docx_path,
    )

    html = paths.html.read_text(encoding='utf-8')
    document = Document(paths.docx)
    paragraph_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)

    assert paths.html == html_path
    assert paths.docx == docx_path
    assert 'Nihai rapor: 2 veri kümesi, 64 geçerleme/seed koşusu' in html
    assert 'CREMA-D – 32 Geçerleme Koşusu' in paragraph_text
    assert 'MELD – 32 Geçerleme Koşusu' in paragraph_text
    assert 'Sınıf-korumalı bootstrap ile held-out test güven aralıkları' in html
    assert '0.3754' in html
    assert '2000 tekrarlı percentile bootstrap' in paragraph_text
    assert 'Held-out Test Olasılık Kalibrasyonu' in paragraph_text
    assert '0.2455' in html
    assert '0.2455 → 0.0575' in html
    assert '0.1049 → 0.0132' in html
    assert 'optimizasyon hedefinin ECE değil NLL olduğu' in paragraph_text
    assert html.count('<img src=') == 8
    assert '96×64 crop-pad macro-F1=0.4707' in paragraph_text
    assert '16 özellik ablation koşusu' in paragraph_text
    assert 'Özellik Temsilinin Çoklu-Seed Doğrulaması' in paragraph_text
    assert 'gradient norm clipping (5.0)' in paragraph_text
    assert '0.4644 ± 0.0063' in html
    assert '0.2370 ± 0.0117' in html
    assert 'eşleştirilmiş 3 seed’in 2 tanesinde' in paragraph_text
    assert 'eşleştirilmiş 3 seed’in 3 tanesinde' in paragraph_text
    assert 'Seed 42 ile yeniden eğitilen 4 aday' in paragraph_text
    assert 'maksimum mutlak fark 0.0000000000' in paragraph_text
    assert '12 çoklu-seed özellik doğrulama koşusu' in paragraph_text
    assert 'Hiperparametrelerin Betimsel Performans İlişkileri' in paragraph_text
    assert 'toplam 60 benzersiz tarama/iyileştirme konfigürasyonunu' in paragraph_text
    assert 'izole nedensel etki kanıtı değil' in paragraph_text
    assert 'n=1 gruplar genellenmemelidir' in paragraph_text
    assert 'temperature scaling' in paragraph_text
    assert 'class="missing"' not in html
    assert html.count('<table>') == 21
    assert len(document.tables) == 21
    assert len(document.inline_shapes) == 8
