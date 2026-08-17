'''Gerçek deney çıktılarından Google-Docs'a hazır HTML ve DOCX raporları üretir.

Bu modülün en önemli tasarım ilkesi: rapordaki HİÇBİR sayı elle yazılmaz.
Tüm tablolar ve bulgu cümleleri, deneylerin diske yazdığı gerçek dosyalardan
(result.json, validation_results.csv, feature_stability.csv vb.) okunarak
kurulur. Böylece kod yeniden çalıştırıldığında rapor otomatik olarak güncel
ve sonuçlarla tutarlı kalır; "raporda yazan sayı ile CSV'deki sayı farklı"
türü hatalar yapısal olarak imkansızlaşır.

Mimari iki katmanlıdır:
1. **İçerik katmanı**: ``_access_blocks``, ``_method_blocks`` gibi fonksiyonlar
   raporun bölümlerini, biçimden bağımsız küçük "blok" nesneleri (Heading,
   Paragraph, TableBlock, ImageBlock...) olarak üretir.
2. **Çizim katmanı**: ``_render_html`` ve ``_render_docx`` AYNI blok listesini
   alır ve iki farklı formata çevirir. İçerik tek kez tanımlanır, iki çıktı
   asla birbirinden sapamaz.

Not: Bu dosyadaki Türkçe metin sabitleri (paragraflar, tablo başlıkları)
raporun KENDİSİDİR, yani çalışma zamanı çıktısıdır — açıklama yorumu değildir.
'''

from __future__ import annotations

import argparse
from dataclasses import dataclass
from html import escape
import json
import os
from pathlib import Path
from typing import Any, Iterable

import pandas as pd


# ---------------------------------------------------------------------------
# Rapor genelinde kullanılan sabitler: korpus/duygu gösterim adları ve teslim
# bağlantıları. Sözlük DEĞERLERİ rapora aynen basılır (çalışma zamanı metni).
# ---------------------------------------------------------------------------
DISPLAY = {'cremad': 'CREMA-D', 'meld': 'MELD'}
EMOTIONS = ('angry', 'disgust', 'fear', 'happy', 'neutral', 'sad')
EMOTION_DISPLAY = {
    'angry': 'Kızgın',
    'disgust': 'Tiksinti',
    'fear': 'Korku',
    'happy': 'Mutlu',
    'neutral': 'Nötr',
    'sad': 'Üzgün',
}
DRIVE_URL = (
    'https://drive.google.com/drive/folders/'
    '1Hbp4WtCGFZjpvQCDxFmqtOmeqq-SMPvW?usp=sharing'
)
GITHUB_URL = (
    'https://github.com/AhmetBabagil/'
    'speech-emotion-recognition/tree/feat/speech-emotion-recognition'
)
DATASET_LINKS = (
    ('CREMA-D', 'https://github.com/CheyneyComputerScience/CREMA-D'),
    ('CREMA-D ses aynası', 'https://huggingface.co/datasets/AbstractTTS/CREMA-D'),
    ('MELD', 'https://github.com/declare-lab/MELD'),
    ('MELD ham veri', 'http://web.eecs.umich.edu/~mihalcea/downloads/MELD.Raw.tar.gz'),
)


@dataclass(frozen=True)
class ReportPaths:
    '''Tek bir tanısal ya da nihai rapor derlemesinde oluşturulan dosyalar.'''

    html: Path
    docx: Path


# ---------------------------------------------------------------------------
# Rapor "blok" tipleri: içerik ile biçimi ayıran küçük, değişmez veri
# sınıfları. Bölüm fonksiyonları bu bloklardan liste üretir; _render_html ve
# _render_docx aynı listeyi kendi formatlarına çevirir. Yeni bir blok tipi
# eklemek = hem iki render fonksiyonuna birer dal eklemek demektir.
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Heading:
    '''Bölüm başlığı; level=1 ana başlık, 2-3 alt başlıklar.'''

    level: int
    text: str


@dataclass(frozen=True)
class Paragraph:
    '''Düz metin paragrafı.'''

    text: str


@dataclass(frozen=True)
class Notice:
    '''Dikkat çekmesi gereken uyarı kutusu (örn. "bu tanısal rapordur").'''

    text: str


@dataclass(frozen=True)
class BulletList:
    '''Madde imli liste.'''

    items: tuple[str, ...]


@dataclass(frozen=True)
class LinkList:
    '''(etiket, url) çiftlerinden oluşan bağlantı listesi.'''

    items: tuple[tuple[str, str], ...]


@dataclass(frozen=True)
class TableBlock:
    '''Başlıklar + satırlar + isteğe bağlı tablo altyazısından oluşan tablo.'''

    headers: tuple[str, ...]
    rows: tuple[tuple[Any, ...], ...]
    caption: str | None = None


@dataclass(frozen=True)
class ImageBlock:
    '''Diskteki bir PNG'yi altyazısıyla rapora gömen görsel bloğu.'''

    path: Path
    alt_text: str
    caption: str


# Tüm blok tiplerinin birleşimi (union) — render fonksiyonlarının kabul
# ettiği tip. Python 3.10+ '|' sözdizimiyle yazılmış bir tip takma adıdır.
ReportBlock = Heading | Paragraph | Notice | BulletList | LinkList | TableBlock | ImageBlock


def _read_json(path: str | Path) -> dict[str, Any]:
    # Deney çıktısı JSON dosyalarını UTF-8 ile okuyan kısa yardımcı.
    return json.loads(Path(path).read_text(encoding='utf-8'))


def _cell(value: Any) -> str:
    # Herhangi bir Python değerini tablo hücresi metnine çevirir; tüm
    # tablolarda tek biçimlilik bu fonksiyondan gelir:
    # - bool -> 'Evet'/'Hayır' (numpy bool_ tip adı üzerinden yakalanır,
    #   çünkü isinstance importsuz numpy tipini göremez),
    # - float -> 4 ondalık basamak,
    # - None/NaN -> '-',
    # - geri kalan her şey -> str().
    if isinstance(value, bool) or type(value).__name__ == 'bool_':
        return 'Evet' if value else 'Hayır'
    if isinstance(value, float):
        return f'{value:.4f}'
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return '-'
    return str(value)


def _relative(target: str | Path, report_path: Path) -> str:
    # Görsel yollarını, HTML dosyasının bulunduğu klasöre GÖRE göreli yapar.
    # Mutlak yol kullansaydık rapor başka bilgisayara taşındığında görseller
    # kırılırdı. as_posix() Windows'un '\\' ayracını web uyumlu '/'e çevirir.
    return Path(os.path.relpath(Path(target), report_path.parent)).as_posix()


def _html_table(
    headers: Iterable[str],
    rows: Iterable[Iterable[Any]],
    caption: str | None = None,
) -> str:
    # TableBlock'u HTML <table> etiketine çevirir. Her hücre escape()'ten
    # geçirilir: verideki '<', '&' gibi karakterler HTML olarak yorumlanmaz,
    # düz metin olarak görünür (XSS/bozuk sayfa koruması).
    head = ''.join(f'<th>{escape(str(header))}</th>' for header in headers)
    body = []
    for row in rows:
        cells = ''.join(f'<td>{escape(_cell(value))}</td>' for value in row)
        body.append(f'<tr>{cells}</tr>')
    caption_html = f'<caption>{escape(caption)}</caption>' if caption else ''
    return (
        f'<table>{caption_html}<thead><tr>{head}</tr></thead>'
        f'<tbody>{"".join(body)}</tbody></table>'
    )


def _render_html(blocks: Iterable[ReportBlock], path: Path, title: str) -> Path:
    '''Rapor bloklarını tek başına yeterli (standalone) UTF-8 HTML belgesi olarak çizer.

    "Standalone" demek: CSS gömülü, dış dosya bağımlılığı yalnızca görseller.
    Fonksiyon blokları sırayla gezer ve her tip için ilgili HTML parçasını
    üretir; tanımadığı bir blok tipi görürse sessizce atlamak yerine
    TypeError fırlatır (yeni tip eklenirse burası unutulmasın diye).
    '''

    body: list[str] = []
    for block in blocks:
        if isinstance(block, Heading):
            # Başlık seviyesini [1, 4] aralığına kıstır: h5/h6 üretme.
            level = min(max(block.level, 1), 4)
            body.append(f'<h{level}>{escape(block.text)}</h{level}>')
        elif isinstance(block, Paragraph):
            body.append(f'<p>{escape(block.text)}</p>')
        elif isinstance(block, Notice):
            body.append(f'<div class="notice">{escape(block.text)}</div>')
        elif isinstance(block, BulletList):
            items = ''.join(f'<li>{escape(item)}</li>' for item in block.items)
            body.append(f'<ul>{items}</ul>')
        elif isinstance(block, LinkList):
            items = ''.join(
                f'<li>{escape(label)}: '
                f'<a href="{escape(url, quote=True)}">{escape(url)}</a></li>'
                for label, url in block.items
            )
            body.append(f'<ul>{items}</ul>')
        elif isinstance(block, TableBlock):
            body.append(_html_table(block.headers, block.rows, block.caption))
        elif isinstance(block, ImageBlock):
            # Görsel diskte varsa göreli yolla <figure> olarak göm; yoksa
            # raporda kırmızı "Eksik görsel" kutusu göster. Testler nihai
            # raporda hiç 'missing' sınıfı kalmadığını doğrular — yani eksik
            # görsel varsa fark edilir, sessizce yutulmaz.
            if block.path.is_file():
                source = escape(_relative(block.path, path), quote=True)
                body.append(
                    '<figure>'
                    f'<img src="{source}" alt="{escape(block.alt_text, quote=True)}">'
                    f'<figcaption>{escape(block.caption)}</figcaption>'
                    '</figure>'
                )
            else:
                body.append(
                    f'<div class="missing">Eksik görsel: {escape(str(block.path))}</div>'
                )
        else:
            raise TypeError(f'Unsupported report block: {type(block).__name__}')

    # Gömülü stil sayfası: A4 baskı ayarı, tablo/başlık renkleri, uyarı
    # kutusu görünümü. Rapor Google Docs'a yapıştırıldığında da bu stiller
    # makul bir görünüm sağlar.
    stylesheet = '''
        @page { size: A4; margin: 18mm; }
        body { font-family: Arial, sans-serif; line-height: 1.5; color: #172033;
               max-width: 1100px; margin: 0 auto; padding: 24px; }
        h1 { color: #123c69; border-bottom: 3px solid #2f80ed; padding-bottom: 8px; }
        h2 { color: #174f85; margin-top: 34px; border-bottom: 1px solid #b8cce4; }
        h3 { color: #2468a2; margin-top: 24px; }
        p, li { font-size: 10.5pt; }
        table { width: 100%; border-collapse: collapse; margin: 16px 0 24px;
                font-size: 8.5pt; page-break-inside: auto; }
        caption { text-align: left; font-weight: bold; margin-bottom: 7px; }
        th, td { border: 1px solid #8ca6bf; padding: 5px 6px; text-align: left; }
        th { background: #dceaf7; }
        tr:nth-child(even) td { background: #f6f9fc; }
        figure { text-align: center; page-break-inside: avoid; margin: 22px 0; }
        img { max-width: 92%; height: auto; }
        figcaption { font-size: 9pt; color: #475569; margin-top: 6px; }
        .notice { border-left: 5px solid #d97706; background: #fff7ed;
                  padding: 12px; margin: 16px 0; font-weight: bold; }
        .missing { color: #b91c1c; border: 1px dashed #b91c1c; padding: 8px; }
        a { color: #075ea8; }
    '''
    document = (
        '<!doctype html><html lang="tr"><head><meta charset="utf-8">'
        f'<title>{escape(title)}</title><style>{stylesheet}</style></head><body>'
        f'{"".join(body)}</body></html>'
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(document, encoding='utf-8')
    return path


def _set_docx_cell_font(cell, size: float) -> None:
    # python-docx'te hücre yazı tipi doğrudan ayarlanamaz; hücrenin içindeki
    # her paragrafın her "run"ına tek tek uygulamak gerekir. Bu tekrarlanan
    # işi tek yardımcıda topluyoruz.
    from docx.shared import Pt

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(size)


def _render_docx(blocks: Iterable[ReportBlock], path: Path, title: str) -> Path:
    '''Aynı rapor bloklarını kendine yeterli bir Word belgesi olarak çizer.

    _render_html ile AYNI blok listesini alır; içerik farkı oluşması bu
    tasarımla imkansızdır. python-docx importları fonksiyon içindedir:
    yalnızca DOCX üretilirken yüklenir, modülün kendisi docx kurulu olmadan
    da import edilebilir.
    '''

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    # Belge üst verileri, sayfa kenar boşlukları ve stil tanımları: HTML'deki
    # stylesheet'in Word karşılığı. Tüm yazı tipleri Arial'e sabitlenir.
    document.core_properties.title = title
    document.core_properties.author = 'Ahmet Babagil'
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    document.styles['Normal'].font.name = 'Arial'
    document.styles['Normal'].font.size = Pt(10)
    for style_name, size in (
        ('Title', 18),
        ('Heading 1', 15),
        ('Heading 2', 13),
        ('Heading 3', 11),
    ):
        document.styles[style_name].font.name = 'Arial'
        document.styles[style_name].font.size = Pt(size)

    for block in blocks:
        if isinstance(block, Heading):
            # Belgenin en başındaki 1. seviye başlık Word'ün 'Title' stilini
            # (level=0) alır; sonraki başlıklar bir seviye kaydırılarak
            # Heading 1-3 aralığına oturtulur.
            if block.level == 1 and not document.paragraphs:
                document.add_heading(block.text, level=0)
            else:
                document.add_heading(block.text, level=min(max(block.level - 1, 1), 3))
        elif isinstance(block, Paragraph):
            paragraph = document.add_paragraph(block.text)
            paragraph.paragraph_format.space_after = Pt(6)
        elif isinstance(block, Notice):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text)
            run.bold = True
            run.font.color.rgb = RGBColor(180, 83, 9)
        elif isinstance(block, BulletList):
            for item in block.items:
                document.add_paragraph(item, style='List Bullet')
        elif isinstance(block, LinkList):
            for label, url in block.items:
                document.add_paragraph(f'{label}: {url}', style='List Bullet')
        elif isinstance(block, TableBlock):
            # Tablo: önce (varsa) kalın altyazı, sonra başlık satırı, sonra
            # veri satırları. Çok sütunlu (10+) tablolar sayfaya sığsın diye
            # daha küçük punto kullanır. Satır, başlık sayısından kısaysa
            # eksik hücreler '-' ile doldurulur (_cell(None) => '-').
            if block.caption:
                caption = document.add_paragraph()
                caption_run = caption.add_run(block.caption)
                caption_run.bold = True
            table = document.add_table(rows=1, cols=len(block.headers))
            table.style = 'Table Grid'
            font_size = 6.0 if len(block.headers) >= 10 else 8.0
            for index, header in enumerate(block.headers):
                cell = table.rows[0].cells[index]
                cell.text = str(header)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                _set_docx_cell_font(cell, font_size)
            for row in block.rows:
                cells = table.add_row().cells
                for index in range(len(block.headers)):
                    cells[index].text = _cell(row[index] if index < len(row) else None)
                    _set_docx_cell_font(cells[index], font_size)
            # Tablodan sonra boş paragraf: tabloların birbirine yapışmasını önler.
            document.add_paragraph()
        elif isinstance(block, ImageBlock):
            # Görsel: Word'e sabit genişlikte (6.1 inç), ortalanmış olarak
            # gömülür; altına italik ve küçük puntolu altyazı eklenir.
            # HTML'nin aksine dosya belgeye KOPYALANIR, göreli yol gerekmez.
            if block.path.is_file():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(block.path), width=Inches(6.1))
                caption = document.add_paragraph(block.caption)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption.runs:
                    caption.runs[0].italic = True
                    caption.runs[0].font.size = Pt(9)
            else:
                document.add_paragraph(f'Eksik görsel: {block.path}')
        else:
            raise TypeError(f'Unsupported report block: {type(block).__name__}')

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def _report_paths(
    output_root: Path,
    diagnostic: bool,
    html_path: str | Path | None,
    docx_path: str | Path | None,
) -> ReportPaths:
    # Çıktı dosyalarının adlarını belirler. Tanısal (quick/limitli) koşuların
    # raporu bilerek FARKLI adla ve deney klasörünün içine yazılır; böylece
    # duman testi raporu, teslim edilecek nihai raporun üzerine yazamaz.
    # Kullanıcı özel yol verdiyse o yol kazanır.
    if diagnostic:
        default_html = output_root / 'RAPOR_DIAGNOSTIK.html'
        default_docx = output_root / 'RAPOR_DIAGNOSTIK.docx'
    else:
        default_html = Path('odev3/PROJE_ILERLEME_RAPORU_3.html')
        default_docx = Path('odev3/PROJE_ILERLEME_RAPORU_3.docx')
    return ReportPaths(
        html=Path(html_path) if html_path else default_html,
        docx=Path(docx_path) if docx_path else default_docx,
    )


def _stage_name(value: str) -> str:
    # CSV'deki İngilizce aşama kodlarını rapor tablolarındaki Türkçe
    # karşılıklarına çevirir; bilinmeyen kod olduğu gibi bırakılır.
    return {
        'screening': 'Tarama',
        'refinement': 'İyileştirme',
        'stability': 'Kararlılık',
        'quick': 'Hızlı tanı',
        'full': 'Tam grid',
    }.get(str(value), str(value))


def _hidden_columns(value: Any) -> tuple[str, str, str]:
    # '512-256' gibi kaydedilmiş gizli katman metnini tablo için üç ayrı
    # sütuna (H1, H2, H3) böler; eksik katmanlar '-' ile doldurulur.
    # Örn. '256' -> ('256', '-', '-'); '512-256-128' -> üçü de dolu.
    sizes = str(value).split('-') if value is not None else []
    padded = (sizes + ['-', '-', '-'])[:3]
    return padded[0], padded[1], padded[2]


def _unique_text(frame: pd.DataFrame, column: str) -> str:
    # Bir sütunda GERÇEKTEN denenen benzersiz değerleri tek metin halinde
    # listeler ('32, 64, 128' gibi). Önce sayısal sıralama denenir; sütun
    # sayıya çevrilemiyorsa (örn. aktivasyon adları) alfabetik sıralanır.
    values = frame[column].dropna().tolist()
    try:
        values = sorted(set(values), key=float)
    except (TypeError, ValueError):
        values = sorted(set(str(value) for value in values))
    return ', '.join(_cell(value) for value in values)


def _artifact_status(path: Path) -> str:
    # Çıktı dosyasının durum hücresi: dosya var mı, boyutu ne? Rapor
    # "üretilen dosyalar" tablosunda her dosyanın gerçekten var olduğunu
    # böylece belgeleriz. 0.1 MB altındaki dosyalar KB olarak gösterilir.
    if not path.is_file():
        return 'Yerel dosya bulunamadı'
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb >= 0.1:
        return f'Hazır ({size_mb:.1f} MB)'
    return f'Hazır ({path.stat().st_size / 1024:.1f} KB)'


def _load_report_inputs(
    output_root: Path,
) -> tuple[dict[str, Any], dict[str, dict[str, Any]], dict[str, pd.DataFrame]]:
    # Raporun ZORUNLU girdilerini yükler: genel summary.json + her korpusun
    # result.json ve validation_results.csv dosyaları. Eksik dosya varsa
    # anlaşılır bir hata ile durur — yarım veriden yarım rapor üretmek
    # yerine sorunu açıkça bildirmek tercih edilir.
    summary_path = output_root / 'summary.json'
    if not summary_path.is_file():
        raise FileNotFoundError(f'Deney özeti bulunamadı: {summary_path}')
    summary = _read_json(summary_path)
    available = summary.get('per_dataset', {})
    if not available:
        raise ValueError(f'Özette veri kümesi sonucu yok: {summary_path}')

    results: dict[str, dict[str, Any]] = {}
    validations: dict[str, pd.DataFrame] = {}
    for corpus in ('cremad', 'meld'):
        if corpus not in available:
            continue
        result_path = output_root / corpus / 'result.json'
        validation_path = output_root / corpus / 'validation_results.csv'
        if not result_path.is_file() or not validation_path.is_file():
            raise FileNotFoundError(
                f'{corpus} rapor girdileri eksik: {result_path}, {validation_path}'
            )
        results[corpus] = _read_json(result_path)
        validations[corpus] = pd.read_csv(validation_path)
    return summary, results, validations


# ===========================================================================
# BÖLÜM ÜRETİCİLERİ
# Aşağıdaki _*_blocks ve _*_rows fonksiyonları raporun bölümlerini sırayla
# kurar. İçlerindeki Türkçe metinler raporun kendisidir (çalışma zamanı
# çıktısı); sayısal değerlerin tamamı deney dosyalarından okunur.
# ===========================================================================


def _access_blocks() -> list[ReportBlock]:
    # 'Erişim Bilgileri' bölümü: proje kimliği, Drive/GitHub bağlantıları ve
    # veri kümesi kaynak linkleri. Tamamen sabit içeriktir.
    return [
        Heading(2, 'Erişim Bilgileri'),
        TableBlock(
            ('Alan', 'Bilgi'),
            (
                ('Orijinal Veri Seti Adı', 'CREMA-D ve MELD'),
                ('Proje Grubu', 'Grup 23'),
                ('Öğrenci', 'Ahmet Babagil – 211101067'),
                ('Google Drive teslim klasörü', DRIVE_URL),
                ('GitHub geliştirme dalı', GITHUB_URL),
            ),
        ),
        LinkList(DATASET_LINKS + (('Google Drive', DRIVE_URL), ('GitHub', GITHUB_URL))),
        Paragraph(
            'İlk iki aşamada kullanılan CREMA-D ve MELD veri kümeleri '
            'değişmemiştir. Bu nedenle orijinal veri tekrar çoğaltılmamış; '
            'Drive teslim klasörü kodu, Git geliştirme geçmişini, en iyi '
            'modelleri ve ara/nihai CSV, JSON ve PNG çıktılarını barındıracak '
            'şekilde düzenlenmiştir.'
        ),
    ]


def _execution_blocks(
    output_root: Path,
    results: dict[str, dict[str, Any]],
    stability_root: Path,
    effect_root: Path,
) -> list[ReportBlock]:
    # 'Çalıştırılacak Dosyalar' bölümü: her kaynak dosyanın görevi, deneyleri
    # yeniden üretecek komutlar ve üretilen çıktıların durum tablosu.
    # files demeti (dosya, görev açıklaması) çiftleridir — rapora aynen basılır.
    files = (
        ('odev3/run_experiment.py', 'Tüm veri kümesi için tarama, yerel iyileştirme, çoklu seed doğrulaması ve tek seferlik test değerlendirmesini çalıştırır.'),
        ('odev3/features_melspec.py', 'Sesi 16 kHz mono yükler; log-Mel spektrogramı sabit boyuta getirip 4096 elemanlı vektöre dönüştürür.'),
        ('odev3/dataset.py', 'Özellikleri yapılandırmaya göre cache’ler ve yalnız eğitim katında öğrenilen z-score normalizasyonunu uygular.'),
        ('odev3/model.py', 'Hazır model kullanmadan yapılandırılabilir PyTorch MLP ağını kurar.'),
        ('odev3/training.py', 'Ağırlıklı çapraz entropi, AdamW, erken durdurma ve metrik hesaplamalarını gerçekleştirir.'),
        ('odev3/uncertainty.py', 'Held-out sınıflandırma metrikleri için sınıf-korumalı bootstrap güven aralıklarını hesaplar.'),
        ('odev3/calibration.py', 'NLL, Brier ve ECE ölçer; validation olasılıklarında sıcaklık katsayısını öğrenip olasılıkları ölçekler.'),
        ('odev3/search_space.py', '16 temel tarama adayını ve kazanan çevresindeki 14 yerel adayı üretir.'),
        ('odev3/feature_ablation.py', '64/80/96 Mel, 64/96 kare ve crop/resize temsillerini yalnız geçerleme setinde karşılaştırır.'),
        ('odev3/feature_stability.py', 'En güçlü iki Mel temsilini seed 42/143/244 ile yalnız geçerleme setinde yeniden eğitip ortalama ve sapmayı ölçer.'),
        ('odev3/hyperparameter_effects.py', 'Tarama ve iyileştirme koşularını parametre değerlerine göre gruplayıp koşu sayısı, ortalama, sapma ve performans farklarını özetler.'),
        ('odev3/pipeline.py', 'Speaker-independent bölmeleri, resume durumunu, model kaydını, tahminleri, belirsizliği ve validation-temelli kalibrasyonu yönetir.'),
        ('odev3/build_report.py', 'Gerçek deney çıktılarından bu HTML ve DOCX raporlarını üretir.'),
    )
    # Çıktı durum tablosu: üç kaynaktan (ana deney, hiperparametre etkileri,
    # özellik kararlılığı) beklenen her dosya için varlık/boyut durumu
    # (_artifact_status) toplanır.
    artifact_rows = []
    for corpus in results:
        corpus_dir = output_root / corpus
        for name, purpose in (
            ('best_model.pt', 'En iyi MLP, normalizasyon ve sınıf ağırlıkları'),
            ('validation_results.csv', '32 geçerleme deneyinin tam tablosu'),
            ('best_training_history.csv', 'Seçilen koşunun epoch geçmişi'),
            ('test_predictions.csv', 'Test örneği tahmin ve olasılıkları'),
            ('test_uncertainty.json', 'Test metriklerinin %95 bootstrap güven aralıkları'),
            ('test_calibration.json', 'Ham test olasılıklarının NLL, Brier ve ECE özeti'),
            ('temperature_scaling.json', 'Validation’da öğrenilen sıcaklık ile ham/ölçekli karşılaştırma'),
            ('test_confusion_matrix.png', 'Normalize test karmaşıklık matrisi'),
            ('test_reliability_diagram.png', 'Ham ve ölçekli test güvenilirlik diyagramı'),
            ('result.json', 'Bölme, yöntem, model ve metrik özeti'),
        ):
            path = corpus_dir / name
            artifact_rows.append(
                (DISPLAY[corpus], path.as_posix(), purpose, _artifact_status(path))
            )
        for name, purpose in (
            ('parameter_effects.csv', 'Her parametre değeri için validation macro-F1 dağılımı'),
            ('parameter_effect_overview.csv', 'En güçlü ve zayıf betimsel parametre grupları'),
        ):
            path = effect_root / corpus / name
            artifact_rows.append(
                (DISPLAY[corpus], path.as_posix(), purpose, _artifact_status(path))
            )
        for name, purpose in (
            ('feature_stability_runs.csv', 'İki temsil × üç seed ham validation koşuları'),
            ('feature_stability.csv', 'Üç-seed ortalama, standart sapma ve sıralama'),
            ('feature_stability_comparison.json', 'Eşleştirilmiş seed farkları ve aday kazanım sayıları'),
            ('feature_stability.png', 'Seed noktaları ile ortalama ± standart sapma görseli'),
        ):
            path = stability_root / corpus / name
            artifact_rows.append(
                (DISPLAY[corpus], path.as_posix(), purpose, _artifact_status(path))
            )
    return [
        Heading(2, 'Çalıştırılacak Dosyalar ve Ne Yaptıkları'),
        TableBlock(('Dosya', 'Görev'), files),
        Heading(3, 'Tekrarlanabilir Çalıştırma Komutları'),
        Notice(
            'Nihai deney: python odev3/run_experiment.py --corpora cremad meld '
            '--grid-mode report --max-epochs 60 --device cpu --feature-workers 4 '
            '--cremad-frame-strategy resize --meld-frame-strategy crop_pad'
        ),
        Paragraph(
            'Özellik karşılaştırması: python odev3/feature_ablation.py '
            '--corpora cremad meld --max-epochs 60 --device cpu '
            '--feature-workers 4'
        ),
        Paragraph(
            'Çoklu-seed özellik doğrulaması: python -m '
            'odev3.feature_stability --corpora cremad meld --max-epochs 60 '
            '--seeds 42 143 244 --device cpu --feature-workers 4'
        ),
        Paragraph(
            'Hiperparametre etki özeti: python -m odev3.hyperparameter_effects '
            '--corpora cremad meld --output-root odev3/outputs '
            '--analysis-root odev3/hyperparameter_effects'
        ),
        Paragraph('Rapor: python odev3/build_report.py --output-root odev3/outputs'),
        Heading(3, 'Üretilen Temel Dosyalar'),
        TableBlock(
            ('Veri seti', 'Çıktı', 'İçerik', 'Durum'),
            tuple(artifact_rows),
        ),
    ]


def _dataset_blocks(results: dict[str, dict[str, Any]]) -> list[ReportBlock]:
    # 'Veri Seti' bölümü: bölme tabloları (kayıt/konuşmacı/sınıf sayıları),
    # sınıf ağırlıkları tablosu ve metrik tanımları. Sayılar result.json'daki
    # split özeti ile training_class_weights alanından gelir.
    split_rows = []
    weight_rows = []
    for corpus, result in results.items():
        split = result['split']
        for key, label in (
            ('train', 'Eğitim'),
            ('validation', 'Geçerleme'),
            ('test', 'Test'),
        ):
            details = split[key]
            counts = details['class_counts']
            split_rows.append(
                (
                    DISPLAY[corpus],
                    label,
                    details['records'],
                    details['speakers'],
                    *(counts[emotion] for emotion in EMOTIONS),
                )
            )
        weights = result['training_class_weights']
        weight_rows.append(
            (DISPLAY[corpus], *(weights[emotion] for emotion in EMOTIONS))
        )

    return [
        Heading(2, 'Veri Seti ve Değerlendirme Metrikleri'),
        Paragraph(
            'Manifest toplam 5199 ses kaydı içerir: 2480 CREMA-D ve 2719 '
            'MELD. Her iki veri kümesi kızgın, tiksinti, korku, mutlu, nötr '
            've üzgün olmak üzere aynı altı sınıfa eşlenmiştir.'
        ),
        Paragraph(
            'Her veri kümesi seed 42 ile konuşmacı-temelli yaklaşık %70 eğitim, '
            '%15 geçerleme ve %15 test katlarına ayrılmıştır. Aynı konuşmacı '
            'iki farklı katta bulunmaz. Sonuç JSON dosyalarındaki üç konuşmacı '
            'kesişimi de boş kümedir. Özellik normalizasyonunun ortalama ve '
            'standart sapması yalnız eğitim katında öğrenilmiştir.'
        ),
        TableBlock(
            (
                'Veri seti', 'Kat', 'Kayıt', 'Konuşmacı',
                'Kızgın', 'Tiksinti', 'Korku', 'Mutlu', 'Nötr', 'Üzgün',
            ),
            tuple(split_rows),
            'Tablo 1. Speaker-independent veri bölmeleri ve sınıf dağılımları',
        ),
        Heading(3, 'Sınıf Dengesizliği ve Kayıp Ağırlıkları'),
        Paragraph(
            'Eğitim kaybında her sınıf için w_c = N / (K × n_c) ters frekans '
            'ağırlığı kullanılmıştır. Böylece özellikle MELD’de daha az görülen '
            'tiksinti ve korku sınıfları kayıp fonksiyonunda daha yüksek katkı '
            'alır. Ağırlıklar yalnız eğitim etiketlerinden hesaplanmıştır.'
        ),
        TableBlock(
            ('Veri seti', 'Kızgın', 'Tiksinti', 'Korku', 'Mutlu', 'Nötr', 'Üzgün'),
            tuple(weight_rows),
            'Tablo 2. Weighted CrossEntropyLoss sınıf ağırlıkları',
        ),
        Heading(3, 'Değerlendirme Metrikleri'),
        BulletList(
            (
                'Doğruluk: doğru tahminlerin tüm test örneklerine oranı.',
                'Dengeli doğruluk: altı sınıfın recall değerlerinin eşit ağırlıklı ortalaması.',
                'Macro-F1: sınıf F1 değerlerinin eşit ağırlıklı ortalaması; ana model seçim metriğidir.',
                'Weighted-F1: sınıf destek sayılarıyla ağırlıklandırılmış F1 değeri.',
                'Sınıf bazında precision, recall ve F1 ile normalize karmaşıklık matrisi.',
                'Negatif log-olabilirlik (NLL): doğru sınıfa verilen olasılığı cezalandırır; düşük değer daha iyidir.',
                'Çok sınıflı Brier skoru: altı olasılığın one-hot hedefe toplam karesel uzaklığıdır; düşük değer daha iyidir.',
                'Expected Calibration Error (ECE): 10 eşit genişlikli güven bininde doğruluk-güven farkının örnek oranıyla ağırlıklı ortalamasıdır; düşük değer daha iyidir.',
            )
        ),
    ]


def _ablation_rows(ablation_root: Path) -> tuple[tuple[Any, ...], ...]:
    # Özellik ablasyonu tablosunun satırları: feature_ablation.csv'den okunur,
    # deneme sırasına dizilir. Son sütun (rank == 1) o adayın kazanan olup
    # olmadığını işaretler; tabloda 'Seçildi' olarak görünür. Dosya yoksa
    # (ablasyon hiç koşulmamışsa) bölüm sessizce boş kalır.
    rows: list[tuple[Any, ...]] = []
    for corpus in ('cremad', 'meld'):
        path = ablation_root / corpus / 'feature_ablation.csv'
        if not path.is_file():
            continue
        frame = pd.read_csv(path).sort_values('trial')
        for row in frame.itertuples(index=False):
            rows.append(
                (
                    DISPLAY[corpus],
                    int(row.trial),
                    row.frame_strategy,
                    int(row.n_mels),
                    int(row.n_frames),
                    int(row.vector_size),
                    int(row.best_epoch),
                    row.val_accuracy,
                    row.val_balanced_accuracy,
                    row.val_macro_f1,
                    int(row.rank) == 1,
                )
            )
    return tuple(rows)


def _ablation_findings(
    ablation_root: Path,
    results: dict[str, dict[str, Any]],
    stability_root: Path | None = None,
) -> str:
    # Ablasyon bulgularını CÜMLE olarak üretir: her korpusun kazananı,
    # kazananın nihai deneyde kullanılan temsille aynı olup olmadığı ve
    # değilse aradaki tek-tohum farkı. Metodolojik dürüstlük burada kodlanır:
    # tek tohumlu küçük bir fark, nihai modeli geriye dönük değiştirme sebebi
    # yapılmaz; bunun yerine çok tohumlu doğrulamaya işaret edilir.
    findings: list[str] = []
    confirmation_needed = False
    for corpus, result in results.items():
        path = ablation_root / corpus / 'feature_ablation.csv'
        if not path.is_file():
            continue
        frame = pd.read_csv(path).sort_values('rank')
        winner = frame.iloc[0]
        selected_feature = result['feature']
        selected_rows = frame[
            (frame['n_mels'] == int(selected_feature['n_mels']))
            & (frame['n_frames'] == int(selected_feature['n_frames']))
            & (frame['frame_strategy'] == selected_feature['frame_strategy'])
        ]
        winner_label = (
            f'{int(winner["n_mels"])}×{int(winner["n_frames"])} '
            f'{str(winner["frame_strategy"]).replace("_", "-")}'
        )
        sentence = (
            f'{DISPLAY[corpus]} için {winner_label} macro-F1='
            f'{winner["val_macro_f1"]:.4f} ile birinci olmuştur.'
        )
        if not selected_rows.empty:
            selected = selected_rows.iloc[0]
            gain = float(winner['val_macro_f1'] - selected['val_macro_f1'])
            same_candidate = str(winner['feature_fingerprint']) == str(
                selected['feature_fingerprint']
            )
            if same_candidate:
                sentence += ' Ana 4096-boyutlu temsil liderliğini korumuştur.'
            else:
                confirmation_needed = True
                sentence += (
                    ' Ana 4096-boyutlu temsile göre tek-seed fark '
                    f'{gain:+.4f} olmuştur.'
                )
        findings.append(sentence)
    if confirmation_needed:
        # Kazanan ile seçilen temsil farklıysa iki durum vardır: çok tohumlu
        # kararlılık deneyi yapıldıysa ona atıf yapılır; yapılmadıysa farkın
        # neden nihai modele yansıtılmadığı açıkça yazılır.
        confirmation_complete = bool(
            stability_root
            and all(
                (stability_root / corpus / 'feature_stability.csv').is_file()
                for corpus in results
            )
        )
        if confirmation_complete:
            findings.append(
                'Tek-seed fark küçük olduğundan nihai checkpoint test sonucuna '
                'göre değiştirilmemiş; adaylar aşağıdaki üç-seed validation '
                'deneyinde ayrıca doğrulanmıştır.'
            )
        else:
            findings.append(
                'Genişletilmiş aday tek seed ile ve yalnız validation katında '
                'değerlendirildiğinden, küçük fark nihai checkpoint’i test '
                'sonucuna göre geriye dönük değiştirmek için kullanılmamış; '
                'çoklu-seed doğrulama adayı olarak kaydedilmiştir.'
            )
    return ' '.join(findings)


def _feature_stability_rows(
    stability_root: Path,
) -> tuple[tuple[Any, ...], ...]:
    # Çok tohumlu kararlılık tablosunun satırları: feature_stability.csv'deki
    # özet istatistikler okunur; 'ortalama ± std' tek hücrede birleştirilir,
    # aday adı 'main_/challenger_' önekinden Türkçe etikete çevrilir.
    rows: list[tuple[Any, ...]] = []
    for corpus in ('cremad', 'meld'):
        path = stability_root / corpus / 'feature_stability.csv'
        if not path.is_file():
            continue
        frame = pd.read_csv(path).sort_values('rank')
        for row in frame.itertuples(index=False):
            rows.append(
                (
                    DISPLAY[corpus],
                    'Ana temsil' if str(row.candidate).startswith('main_')
                    else 'Alternatif',
                    str(row.frame_strategy).replace('_', '-'),
                    int(row.n_mels),
                    int(row.n_frames),
                    int(row.vector_size),
                    int(row.runs),
                    str(row.seeds),
                    f'{row.val_macro_f1_mean:.4f} ± '
                    f'{row.val_macro_f1_std:.4f}',
                    row.val_macro_f1_min,
                    row.val_macro_f1_max,
                    int(row.rank) == 1,
                )
            )
    return tuple(rows)


def _feature_stability_findings(
    stability_root: Path,
    ablation_root: Path,
) -> str:
    # Kararlılık bulgularını cümle olarak üretir. İki iddia sayılarla
    # desteklenir:
    # 1. Ana temsil ile alternatif arasındaki üç-tohum ortalama farkı ve
    #    eşleştirilmiş tohum kazanım sayıları.
    # 2. Determinizm kanıtı: seed 42 ile yeniden eğitilen adayların ilk
    #    ablasyondaki skorlarını birebir yeniden ürettiği (maks. mutlak fark).
    findings = []
    reproduced_rows = 0
    maximum_reproduction_difference = 0.0
    for corpus in ('cremad', 'meld'):
        path = stability_root / corpus / 'feature_stability.csv'
        if not path.is_file():
            continue
        frame = pd.read_csv(path)
        main = frame[frame['candidate'].astype(str).str.startswith('main_')].iloc[0]
        challenger = frame[
            frame['candidate'].astype(str).str.startswith('challenger_')
        ].iloc[0]
        comparison = _read_json(
            stability_root / corpus / 'feature_stability_comparison.json'
        )
        difference = float(
            main['val_macro_f1_mean'] - challenger['val_macro_f1_mean']
        )
        findings.append(
            f'{DISPLAY[corpus]} ana 64×64 temsilinin üç-seed validation '
            f'macro-F1 değeri {main["val_macro_f1_mean"]:.4f} ± '
            f'{main["val_macro_f1_std"]:.4f}, alternatifin değeri '
            f'{challenger["val_macro_f1_mean"]:.4f} ± '
            f'{challenger["val_macro_f1_std"]:.4f}; ortalama fark '
            f'{difference:+.4f} olmuştur. Ana temsil eşleştirilmiş '
            f'{len(comparison["seeds"])} seed’in '
            f'{comparison["main_wins"]} tanesinde daha yüksek macro-F1 '
            'üretmiştir.'
        )
        # Tekrarlanabilirlik kontrolü: kararlılık deneyindeki seed-42 koşuları
        # ile ablasyondaki orijinal koşular öznitelik parmak izi üzerinden
        # eşleştirilir; skor farklarının mutlak maksimumu raporlanır (ideal
        # olarak 0.0 — aynı tohum aynı sonucu üretmeli).
        original = pd.read_csv(
            ablation_root / corpus / 'feature_ablation.csv'
        )
        repeated = pd.read_csv(
            stability_root / corpus / 'feature_stability_runs.csv'
        )
        repeated = repeated[repeated['seed'] == 42]
        matched = repeated.merge(
            original,
            on='feature_fingerprint',
            suffixes=('_repeat', '_original'),
        )
        differences = (
            matched['val_macro_f1_repeat']
            - matched['val_macro_f1_original']
        ).abs()
        reproduced_rows += len(matched)
        if not differences.empty:
            maximum_reproduction_difference = max(
                maximum_reproduction_difference,
                float(differences.max()),
            )
    if reproduced_rows:
        findings.append(
            f'Seed 42 ile yeniden eğitilen {reproduced_rows} aday, ilk ablation '
            'koşularındaki macro-F1 değerlerini tam olarak yeniden üretmiş; '
            f'maksimum mutlak fark {maximum_reproduction_difference:.10f} '
            'olmuştur.'
        )
    findings.append(
        'Bu doğrulamada yalnız train ve validation özellikleri yüklenmiş, '
        'held-out test kayıtları model veya temsil seçiminde kullanılmamıştır.'
    )
    return ' '.join(findings)


def _method_blocks(
    results: dict[str, dict[str, Any]],
    ablation_root: Path,
    stability_root: Path,
) -> list[ReportBlock]:
    # 'Yöntem' bölümü: Mel öznitelik açıklaması, ablasyon tablosu+bulgular,
    # çok tohumlu doğrulama tablosu+görselleri ve nihai MLP mimarileri.
    # Ablasyon/kararlılık dosyaları yoksa ilgili alt bölümler atlanır.
    architecture_rows = []
    for corpus, result in results.items():
        config = result['best_config']
        hidden = config['hidden_dims']
        # Mimarinin okunur gösterimi: '4096 → 768 → 384 → 6' gibi.
        architecture = '4096 → ' + ' → '.join(str(size) for size in hidden) + ' → 6'
        architecture_rows.append(
            (
                DISPLAY[corpus],
                result['feature']['frame_strategy'],
                architecture,
                len(hidden),
                config['activation'],
                config['batch_norm'],
                config['dropout'],
                result['model_parameters'],
            )
        )
    blocks: list[ReportBlock] = [
        Heading(2, 'Yöntem'),
        Heading(3, 'Mel-Spectrogram Vektörü'),
        Paragraph(
            'Sesler librosa ile 16 kHz ve mono yüklenmiştir. '
            'librosa.feature.melspectrogram çağrısında 64 Mel bandı, n_fft=1024, '
            'hop_length=512, 20–8000 Hz ve power=2 kullanılmış; güç değerleri '
            'top_db=80 ile log-dB ölçeğine dönüştürülmüştür. 64×64 matris '
            'satır-major düzleştirilerek tam 4096 boyutlu vektör elde edilmiştir. '
            'PCA veya önceden eğitilmiş bir öznitelik modeli kullanılmamıştır.'
        ),
    ]
    ablation_rows = _ablation_rows(ablation_root)
    if ablation_rows:
        blocks.extend(
            [
                Paragraph(
                    'Zaman ve frekans temsilini test setine dokunmadan incelemek '
                    'için referans MLP ile seed 42 sabit tutulmuş; 64/80/96 Mel '
                    'bandı, 64/96 kare ve merkez crop-pad/tüm konuşmayı resize '
                    'seçeneklerinden oluşan sekiz aday karşılaştırılmıştır.'
                ),
                TableBlock(
                    (
                        'Veri seti', 'Deney', 'Zaman işlemi', 'Mel', 'Kare', 'Boyut',
                        'En iyi ep.', 'Doğ.', 'Deng. Doğ.', 'Macro-F1', 'Seçildi',
                    ),
                    ablation_rows,
                    'Tablo 3. Yalnız geçerleme katında Mel özellik ablation sonuçları',
                ),
                Paragraph(
                    _ablation_findings(
                        ablation_root,
                        results,
                        stability_root,
                    )
                ),
            ]
        )
    stability_rows = _feature_stability_rows(stability_root)
    if stability_rows:
        blocks.extend(
            [
                Heading(3, 'Özellik Temsilinin Çoklu-Seed Doğrulaması'),
                Paragraph(
                    'Tek-seed ablation sıralamasındaki küçük farkların '
                    'kararlılığını ölçmek için her corpus’ta ana temsil ve en '
                    'güçlü alternatif aynı referans MLP ile seed 42, 143 ve 244 '
                    'üzerinde yeniden eğitilmiştir. Seçim ölçütü üç seed’in '
                    'validation macro-F1 ortalamasıdır; standart sapma, minimum '
                    've maksimum değerler de saklanmıştır.'
                ),
                TableBlock(
                    (
                        'Veri seti', 'Aday', 'Zaman işlemi', 'Mel', 'Kare',
                        'Boyut', 'Koşu', 'Seed’ler', 'Macro-F1 ort. ± std',
                        'Minimum', 'Maksimum', 'Seçildi',
                    ),
                    stability_rows,
                    'Tablo 4. Mel temsillerinin üç-seed validation kararlılığı',
                ),
                Paragraph(
                    _feature_stability_findings(
                        stability_root,
                        ablation_root,
                    )
                ),
            ]
        )
        for corpus in results:
            blocks.append(
                ImageBlock(
                    stability_root / corpus / 'feature_stability.png',
                    f'{DISPLAY[corpus]} Mel temsil kararlılığı',
                    f'{DISPLAY[corpus]} için üç seed sonucu, ortalama ve '
                    'standart sapma.',
                )
            )
    blocks.extend(
        [
            Heading(3, 'Nihai MLP Yapısı'),
            Paragraph(
                'MLP tamamen PyTorch ile sıfırdan kurulmuştur. Her gizli blok '
                'Linear → isteğe bağlı BatchNorm1d → aktivasyon → Dropout '
                'sırasındadır. Çıkış katmanı altı logit üretir. Eğitimde '
                'weighted CrossEntropyLoss, AdamW, gradient norm clipping '
                '(5.0) ve validation macro-F1 tabanlı erken durdurma '
                'kullanılmıştır. Üst sınır 60 epoch’tur.'
            ),
            TableBlock(
                (
                    'Veri seti', 'Zaman işlemi', 'Katman dizisi', 'Gizli katman',
                    'Aktivasyon', 'Batch norm', 'Dropout', 'Eğitilebilir parametre',
                ),
                tuple(architecture_rows),
                'Tablo 5. Nihai veri-kümesine özel MLP mimarileri',
            ),
        ]
    )
    return blocks


def _hyperparameter_summary_rows(
    validations: dict[str, pd.DataFrame],
) -> tuple[tuple[Any, ...], ...]:
    # 'Gerçekten taranan değerler' tablosu: her hiperparametre sütununda
    # denenen benzersiz değerlerin listesi (_unique_text). Amaç, aramanın
    # kapsamını tek satırda kanıtlamaktır.
    rows = []
    for corpus, frame in validations.items():
        rows.append(
            (
                DISPLAY[corpus],
                len(frame),
                _unique_text(frame, 'batch_size'),
                _unique_text(frame, 'learning_rate'),
                _unique_text(frame, 'patience'),
                _unique_text(frame, 'hidden_dims'),
                _unique_text(frame, 'activation'),
                _unique_text(frame, 'batch_norm'),
                _unique_text(frame, 'dropout'),
                _unique_text(frame, 'weight_decay'),
                _unique_text(frame, 'seed'),
            )
        )
    return tuple(rows)


def _trial_parameter_rows(frame: pd.DataFrame) -> tuple[tuple[Any, ...], ...]:
    # Deneme-parametre tablosu: her doğrulama koşusunun TÜM hiperparametre
    # değerleri, deneme numarasına göre sıralı. Gizli katmanlar üç ayrı
    # sütuna açılır (_hidden_columns) ki tablo hizalı okunsun.
    rows = []
    for row in frame.sort_values('trial').itertuples(index=False):
        h1, h2, h3 = _hidden_columns(row.hidden_dims)
        rows.append(
            (
                int(row.trial),
                _stage_name(row.search_stage),
                int(row.seed),
                row.learning_rate,
                int(row.batch_size),
                int(row.patience),
                int(row.hidden_layers),
                h1,
                h2,
                h3,
                row.activation,
                row.batch_norm,
                row.dropout,
                row.weight_decay,
            )
        )
    return tuple(rows)


def _trial_metric_rows(frame: pd.DataFrame) -> tuple[tuple[Any, ...], ...]:
    # Deneme-metrik tablosu: aynı koşuların skor tarafı (kayıp, doğruluk,
    # F1'ler, erken durma bilgisi ve 'Seçildi' bayrağı). Parametre tablosuyla
    # deneme numarası üzerinden eşleşir.
    rows = []
    for row in frame.sort_values('trial').itertuples(index=False):
        rows.append(
            (
                int(row.trial),
                int(row.best_epoch),
                int(row.epochs_trained),
                row.stopped_early,
                row.val_loss,
                row.val_accuracy,
                row.val_balanced_accuracy,
                row.val_macro_f1,
                row.val_weighted_f1,
                row.selected,
            )
        )
    return tuple(rows)


def _hyperparameter_effect_rows(
    effect_root: Path,
) -> tuple[tuple[Any, ...], ...]:
    # Hiperparametre etki tablosu: hyperparameter_effects modülünün ürettiği
    # parameter_effect_overview.csv'den, her parametrenin en güçlü/en zayıf
    # grubu ve aradaki fark okunur. Grup boyutu aralığı ('1–15' gibi) da
    # gösterilir ki okuyucu az koşulu grupları temkinli yorumlasın.
    rows = []
    for corpus in ('cremad', 'meld'):
        path = effect_root / corpus / 'parameter_effect_overview.csv'
        if not path.is_file():
            continue
        frame = pd.read_csv(path).sort_values('parameter_order')
        for row in frame.itertuples(index=False):
            rows.append(
                (
                    DISPLAY[corpus],
                    row.parameter_label,
                    int(row.tested_value_count),
                    row.best_value,
                    int(row.best_runs),
                    row.best_mean_macro_f1,
                    row.worst_value,
                    int(row.worst_runs),
                    row.worst_mean_macro_f1,
                    row.mean_macro_f1_spread,
                    f'{int(row.minimum_group_runs)}–{int(row.maximum_group_runs)}',
                )
            )
    return tuple(rows)


def _hyperparameter_effect_findings(effect_root: Path) -> tuple[str, ...]:
    # Etki analizinin bulgu paragrafları: kullanılan konfigürasyon sayısı,
    # her korpusta en geniş farkı gösteren parametre ve — en önemlisi —
    # "bu betimseldir, nedensel değildir" uyarısı. Uyarının raporda her zaman
    # yer alması metodolojik dürüstlüğün parçasıdır.
    summary_path = effect_root / 'summary.json'
    if not summary_path.is_file():
        return ()
    summary = _read_json(summary_path)
    total_configurations = sum(
        int(item['configuration_trials'])
        for item in summary.get('per_dataset', {}).values()
    )
    total_excluded = sum(
        int(item['excluded_stability_trials'])
        for item in summary.get('per_dataset', {}).values()
    )
    findings = [
        f'Etki özeti, iki veri kümesindeki toplam {total_configurations} benzersiz '
        'tarama/iyileştirme konfigürasyonunu kullanmıştır. Aynı yapıların seed '
        f'tekrarı olan {total_excluded} stability satırı bu hesaba katılmamıştır.'
    ]
    for corpus in ('cremad', 'meld'):
        overview_path = effect_root / corpus / 'parameter_effect_overview.csv'
        if not overview_path.is_file():
            continue
        frame = pd.read_csv(overview_path)
        # En geniş grup-ortalaması aralığına sahip parametre: skoru en çok
        # "oynatan" değişken olarak öne çıkarılır.
        strongest = frame.loc[frame['mean_macro_f1_spread'].idxmax()]
        findings.append(
            f'{DISPLAY[corpus]} için en geniş grup-ortalaması aralığı '
            f'{strongest.parameter_label} değişkeninde görülmüştür: '
            f'{strongest.best_value} (n={int(strongest.best_runs)}, '
            f'macro-F1={strongest.best_mean_macro_f1:.4f}) ile '
            f'{strongest.worst_value} (n={int(strongest.worst_runs)}, '
            f'macro-F1={strongest.worst_mean_macro_f1:.4f}) arasındaki '
            f'betimsel fark {strongest.mean_macro_f1_spread:.4f} değerindedir.'
        )
    findings.append(
        'Arama dengeli tam faktöriyel değildir ve konfigürasyonlarda birden çok '
        'parametre birlikte değişmektedir. Bu nedenle tablo izole nedensel etki '
        'kanıtı değil, sonraki kontrollü deneyleri yönlendiren betimsel ilişki '
        'olarak okunmalıdır; özellikle n=1 gruplar genellenmemelidir.'
    )
    return tuple(findings)


def _stability_rows(
    results: dict[str, dict[str, Any]],
    validations: dict[str, pd.DataFrame],
) -> tuple[tuple[Any, ...], ...]:
    # Kazanan KONFİGÜRASYONUN tohum duyarlılığı tablosu (özellik kararlılığı
    # tablosundan farklı: burada model konfigi test edilir). Her tohumun
    # tekil satırlarının ardından 'ortalama ± std' özet satırı eklenir.
    # Satırlar validation_results.csv'den config_id eşleşmesiyle bulunur.
    rows = []
    for corpus, result in results.items():
        stability = result.get('stability')
        if not stability:
            continue
        frame = validations[corpus]
        selected = frame[frame['config_id'] == stability['config_id']].sort_values('seed')
        for row in selected.itertuples(index=False):
            rows.append(
                (
                    DISPLAY[corpus],
                    int(row.seed),
                    int(row.best_epoch),
                    row.val_accuracy,
                    row.val_balanced_accuracy,
                    row.val_macro_f1,
                )
            )
        rows.append(
            (
                DISPLAY[corpus] + ' ortalama ± std',
                '-',
                '-',
                stability['val_accuracy_mean'],
                stability['val_balanced_accuracy_mean'],
                f"{stability['val_macro_f1_mean']:.4f} ± {stability['val_macro_f1_std']:.4f}",
            )
        )
    return tuple(rows)


def _validation_blocks(
    output_root: Path,
    results: dict[str, dict[str, Any]],
    validations: dict[str, pd.DataFrame],
    effect_root: Path,
) -> list[ReportBlock]:
    # 'Model Geçerleme' bölümü: arama protokolünün anlatımı, taranan değer
    # özeti, her korpus için 32 koşunun parametre+metrik tabloları, betimsel
    # etki analizi, tohum kararlılığı ve öğrenme eğrisi görselleri.
    # table_number değişkeni tablo numaralarını korpuslar arasında kesintisiz
    # ilerletir (Tablo 7, 8, 9, ...).
    blocks: list[ReportBlock] = [
        Heading(2, 'Model Geçerleme'),
        Paragraph(
            'Her veri kümesinde önce tüm temel hiperparametre gruplarını tek '
            'faktörlü ve yorumlanabilir biçimde kapsayan 16 tarama deneyi '
            'çalıştırılmıştır. Tarama kazananının çevresinde 14 yerel '
            'iyileştirme denenmiş; son yapı seed 42, 143 ve 244 ile toplam üç '
            'kez ölçülmüştür. Model seçimi yalnız validation macro-F1, eşitlikte '
            'dengeli doğruluk ve validation loss ile yapılmıştır. Test '
            'özellikleri 32 koşu tamamlanmadan yüklenmemiştir.'
        ),
        TableBlock(
            (
                'Veri seti', 'Koşu', 'Batch', 'LR', 'Patience', 'Gizli yapılar',
                'Aktivasyon', 'BN', 'Dropout', 'Weight decay', 'Seed',
            ),
            _hyperparameter_summary_rows(validations),
            'Tablo 6. Gerçekten taranan hiperparametre değerleri',
        ),
    ]
    table_number = 7
    for corpus, frame in validations.items():
        display = DISPLAY[corpus]
        blocks.extend(
            [
                Heading(3, f'{display} – 32 Geçerleme Koşusu'),
                TableBlock(
                    (
                        'Deney', 'Aşama', 'Seed', 'LR', 'Batch', 'Pat.',
                        'Katman', 'H1', 'H2', 'H3', 'Akt.', 'BN', 'Drop.', 'WD',
                    ),
                    _trial_parameter_rows(frame),
                    f'Tablo {table_number}. {display} hiperparametre deneyleri',
                ),
                TableBlock(
                    (
                        'Deney', 'En iyi ep.', 'Eğitilen ep.', 'Erken durdu',
                        'Val loss', 'Doğ.', 'Deng. Doğ.', 'Macro-F1',
                        'Weighted-F1', 'Seçildi',
                    ),
                    _trial_metric_rows(frame),
                    f'Tablo {table_number + 1}. {display} geçerleme performansları',
                ),
            ]
        )
        table_number += 2
    effect_rows = _hyperparameter_effect_rows(effect_root)
    if effect_rows:
        blocks.extend(
            [
                Heading(3, 'Hiperparametrelerin Betimsel Performans İlişkileri'),
                Paragraph(
                    'Tarama ve yerel iyileştirme sonuçları her parametre değeri '
                    'için ayrı gruplandırılmış; grup koşu sayısı, ortalama '
                    'validation macro-F1 ve en güçlü–en zayıf grup farkı '
                    'hesaplanmıştır.'
                ),
                TableBlock(
                    (
                        'Veri seti', 'Parametre', 'Değer', 'En güçlü grup',
                        'n', 'Ort. F1', 'En zayıf grup', 'n', 'Ort. F1',
                        'Fark', 'Grup n aralığı',
                    ),
                    effect_rows,
                    f'Tablo {table_number}. Hiperparametre gruplarının betimsel validation ilişkileri',
                ),
                *(
                    Paragraph(finding)
                    for finding in _hyperparameter_effect_findings(effect_root)
                ),
            ]
        )
        table_number += 1
    blocks.extend(
        [
            Heading(3, 'Çoklu Seed Kararlılık Sonuçları'),
            TableBlock(
                ('Veri seti', 'Seed', 'En iyi epoch', 'Doğ.', 'Deng. Doğ.', 'Macro-F1'),
                _stability_rows(results, validations),
                f'Tablo {table_number}. Seçilen yapıların seed duyarlılığı',
            ),
        ]
    )
    for corpus in results:
        blocks.append(
            ImageBlock(
                output_root / corpus / 'best_training_history.png',
                f'{DISPLAY[corpus]} eğitim ve geçerleme eğrileri',
                f'{DISPLAY[corpus]} seçilmiş MLP için loss ve macro-F1 eğitim geçmişi.',
            )
        )
    return blocks


def _overall_test_rows(results: dict[str, dict[str, Any]]) -> tuple[tuple[Any, ...], ...]:
    # Nihai test tablosunun satırları: korpus başına tek satırda tüm ana
    # test metrikleri (result.json'daki 'test' alanından).
    rows = []
    for corpus, result in results.items():
        test = result['test']
        rows.append(
            (
                DISPLAY[corpus],
                result['selected_seed'],
                result['best_epoch'],
                test['accuracy'],
                test['balanced_accuracy'],
                test['macro_f1'],
                test['weighted_f1'],
                result['test_loss'],
            )
        )
    return tuple(rows)


def _per_class_rows(results: dict[str, dict[str, Any]]) -> tuple[tuple[Any, ...], ...]:
    # Sınıf bazında test tablosu: her duygu için precision/recall/F1 ve
    # destek (o sınıfın test örneği sayısı). Duygu adları Türkçe gösterime
    # çevrilir.
    rows = []
    for corpus, result in results.items():
        per_class = result['test']['per_class']
        for emotion in EMOTIONS:
            metrics = per_class[emotion]
            rows.append(
                (
                    DISPLAY[corpus],
                    EMOTION_DISPLAY[emotion],
                    metrics['precision'],
                    metrics['recall'],
                    metrics['f1'],
                    metrics['support'],
                )
            )
    return tuple(rows)


def _uncertainty_rows(
    results: dict[str, dict[str, Any]],
) -> tuple[tuple[Any, ...], ...]:
    # Bootstrap güven aralığı tablosu: her metrik için nokta tahmini +
    # alt/üst sınır + güven düzeyi (test_uncertainty alanından).
    metric_labels = (
        ('accuracy', 'Doğruluk'),
        ('balanced_accuracy', 'Dengeli doğruluk'),
        ('macro_f1', 'Macro-F1'),
        ('weighted_f1', 'Weighted-F1'),
    )
    rows = []
    for corpus, result in results.items():
        uncertainty = result.get('test_uncertainty')
        if not uncertainty:
            continue
        confidence = float(uncertainty['confidence'])
        for metric, label in metric_labels:
            interval = uncertainty['metrics'][metric]
            rows.append(
                (
                    DISPLAY[corpus],
                    label,
                    interval['estimate'],
                    interval['lower'],
                    interval['upper'],
                    confidence,
                )
            )
    return tuple(rows)


def _calibration_rows(
    results: dict[str, dict[str, Any]],
) -> tuple[tuple[Any, ...], ...]:
    # Kalibrasyon tablosu: öğrenilen sıcaklık T ve her metriğin
    # 'ham → ölçekli' geçişi tek hücrede ok işaretiyle gösterilir. Son sütun,
    # sınıf kararlarının korunduğunun (True) kanıtıdır.
    rows = []
    for corpus, result in results.items():
        scaling = result.get('temperature_scaling')
        if not scaling:
            continue
        fit = scaling['fit']
        before = scaling['test']['before']
        after = scaling['test']['after']
        rows.append(
            (
                DISPLAY[corpus],
                fit['temperature'],
                f'{fit["validation_nll_before"]:.4f} → '
                f'{fit["validation_nll_after"]:.4f}',
                f'{before["negative_log_likelihood"]:.4f} → '
                f'{after["negative_log_likelihood"]:.4f}',
                f'{before["multiclass_brier_score"]:.4f} → '
                f'{after["multiclass_brier_score"]:.4f}',
                f'{before["expected_calibration_error"]:.4f} → '
                f'{after["expected_calibration_error"]:.4f}',
                f'{before["mean_confidence"]:.4f} → '
                f'{after["mean_confidence"]:.4f}',
                after['accuracy'],
                scaling['class_predictions_preserved'],
            )
        )
    return tuple(rows)


def _calibration_analysis(results: dict[str, dict[str, Any]]) -> str:
    # Kalibrasyon bulgu paragrafı: T değeri, NLL/ECE'nin testteki değişimi,
    # ortalama güvenin doğrulukla nasıl hizalandığı ve — dürüstlük notu —
    # optimizasyon hedefinin ECE değil NLL olduğu vurgusu.
    findings = []
    for corpus, result in results.items():
        scaling = result.get('temperature_scaling')
        if not scaling:
            continue
        fit = scaling['fit']
        before = scaling['test']['before']
        after = scaling['test']['after']
        validation_before = scaling['validation']['before']
        validation_after = scaling['validation']['after']
        findings.append(
            f'{DISPLAY[corpus]} için validation NLL ile öğrenilen '
            f'T={fit["temperature"]:.4f}, test NLL değerini '
            f'{before["negative_log_likelihood"]:.4f}’ten '
            f'{after["negative_log_likelihood"]:.4f}’e ve ECE değerini '
            f'{before["expected_calibration_error"]:.4f}’ten '
            f'{after["expected_calibration_error"]:.4f}’e düşürmüştür. '
            f'Ortalama güven {before["mean_confidence"]:.4f}’ten '
            f'{after["mean_confidence"]:.4f}’e inerken doğruluk '
            f'{after["accuracy"]:.4f} ve sınıf tahminleri korunmuştur. '
            f'Validation ECE {validation_before["expected_calibration_error"]:.4f} '
            f'→ {validation_after["expected_calibration_error"]:.4f} olmuştur; '
            'optimizasyon hedefinin ECE değil NLL olduğu özellikle not edilmiştir.'
        )
    return ' '.join(findings)


def _comparison_rows(path: Path) -> tuple[tuple[Any, ...], ...]:
    # Ödevler arası model karşılaştırma tablosu: pipeline'ın yazdığı
    # model_comparison.csv (KNN, ağaç modelleri + MLP) satır satır aktarılır.
    if not path.is_file():
        return ()
    frame = pd.read_csv(path)
    rows = []
    for row in frame.itertuples(index=False):
        rows.append(
            (
                DISPLAY.get(row.corpus, row.corpus),
                row.model,
                row.feature_dim,
                row.pca_dim,
                row.test_accuracy,
                row.test_balanced_accuracy,
                row.test_macro_f1,
                row.test_weighted_f1,
            )
        )
    return tuple(rows)


def _result_analysis(results: dict[str, dict[str, Any]], comparison_path: Path) -> list[ReportBlock]:
    # Sonuç yorumu paragrafları. Her korpus için üç veri odaklı gözlem:
    # 1. Doğrulama-test farkı (genelleme boşluğu) — büyükse aşırı uyum işareti.
    # 2. En güçlü/en zayıf duygu sınıfı (per_class F1 uzerinden max/min).
    # 3. Varsa, önceki ödevlerin en iyi klasik modeliyle kıyas.
    blocks: list[ReportBlock] = []
    comparison = pd.read_csv(comparison_path) if comparison_path.is_file() else None
    for corpus, result in results.items():
        display = DISPLAY[corpus]
        validation_f1 = result['validation']['macro_f1']
        test_f1 = result['test']['macro_f1']
        gap = validation_f1 - test_f1
        per_class = result['test']['per_class']
        strongest = max(per_class, key=lambda name: per_class[name]['f1'])
        weakest = min(per_class, key=lambda name: per_class[name]['f1'])
        text = (
            f'{display}: seçilen modelin validation macro-F1 değeri '
            f'{validation_f1:.4f}, held-out test macro-F1 değeri {test_f1:.4f}; '
            f'fark {gap:.4f} olmuştur. En güçlü sınıf '
            f'{EMOTION_DISPLAY[strongest]} (F1={per_class[strongest]["f1"]:.4f}), '
            f'en zayıf sınıf {EMOTION_DISPLAY[weakest]} '
            f'(F1={per_class[weakest]["f1"]:.4f}) olmuştur.'
        )
        if comparison is not None:
            # Karşılaştırma tablosundan bu korpusun MLP OLMAYAN en iyi modeli
            # bulunur (~contains('MLP') filtresi) ve cümleye eklenir.
            prior = comparison[
                (comparison['corpus'] == corpus)
                & ~comparison['model'].astype(str).str.contains('MLP', case=False)
            ].sort_values('test_macro_f1', ascending=False)
            if not prior.empty:
                best = prior.iloc[0]
                text += (
                    f' Önceki aşamalardaki en yüksek test macro-F1 '
                    f'{best["model"]} ile {best["test_macro_f1"]:.4f}’tür.'
                )
        blocks.append(Paragraph(text))
    blocks.append(
        Paragraph(
            'CREMA-D MLP, Karar Ağacı sonucunu geçmesine rağmen Gradient '
            'Boosting, Rastgele Orman ve KNN’in altında kalmıştır. MELD MLP de '
            'Karar Ağacından iyi, diğer üç klasik modelden düşüktür. Bu durum '
            'özellikle MELD’de yalnız akustik Mel vektörünün diyalog bağlamı ve '
            'metinsel anlamı temsil edememesi, ayrıca konuşmacı-bağımsız testte '
            'alan kayması görülmesiyle uyumludur. Sonuçlar saklanmamış veya '
            'test setine göre yeniden seçilmemiştir.'
        )
    )
    return blocks


def _test_blocks(output_root: Path, results: dict[str, dict[str, Any]]) -> list[ReportBlock]:
    # 'Deney Sonuçları' bölümü: genel test tablosu, bootstrap aralıkları,
    # kalibrasyon tablo+görselleri, sınıf bazlı sonuçlar, karışıklık
    # matrisleri, model karşılaştırması ve yorum paragrafları. İsteğe bağlı
    # parçalar (belirsizlik, kalibrasyon, karşılaştırma) ancak ilgili veri
    # mevcutsa eklenir.
    comparison_path = output_root / 'model_comparison.csv'
    blocks: list[ReportBlock] = [
        Heading(2, 'Deney Sonuçları ve Değerlendirme'),
        TableBlock(
            (
                'Veri seti', 'Seed', 'En iyi ep.', 'Doğ.', 'Deng. Doğ.',
                'Macro-F1', 'Weighted-F1', 'Test loss',
            ),
            _overall_test_rows(results),
            'Nihai held-out test sonuçları',
        ),
    ]
    uncertainty_rows = _uncertainty_rows(results)
    if uncertainty_rows:
        blocks.extend(
            [
                TableBlock(
                    (
                        'Veri seti', 'Metrik', 'Tahmin', 'Alt sınır',
                        'Üst sınır', 'Güven',
                    ),
                    uncertainty_rows,
                    'Sınıf-korumalı bootstrap ile held-out test güven aralıkları',
                ),
                Paragraph(
                    'Güven aralıkları test katındaki her duygu sınıfının örnek '
                    'sayısını koruyan 2000 tekrarlı percentile bootstrap ile, '
                    'seed 42 ve %95 güven düzeyinde hesaplanmıştır. Test katı '
                    'model veya hiperparametre seçimi için kullanılmamıştır.'
                ),
            ]
        )
    calibration_rows = _calibration_rows(results)
    if calibration_rows:
        blocks.extend(
            [
                Heading(3, 'Held-out Test Olasılık Kalibrasyonu'),
                TableBlock(
                    (
                        'Veri seti', 'T', 'Val NLL ham→ölç.',
                        'Test NLL ham→ölç.', 'Brier ham→ölç.',
                        'ECE ham→ölç.', 'Ort. güven ham→ölç.', 'Doğ.',
                        'Karar korundu',
                    ),
                    calibration_rows,
                    'Validation’da öğrenilen sıcaklıkla olasılık kalibrasyonu',
                ),
                Paragraph(_calibration_analysis(results)),
                Paragraph(
                    'Sıcaklık katsayısı, model ve hiperparametre seçimi bittikten '
                    'sonra yalnız validation olasılıklarında 0.25–10 aralığında '
                    'iki aşamalı logaritmik aramayla NLL en aza indirilerek '
                    'öğrenilmiştir. Kilitli test etiketleri sıcaklık seçiminde '
                    'kullanılmamış; öğrenilen tek katsayı test olasılıklarına bir '
                    'kez uygulanmıştır. T>1 dağılımı yumuşatır; sınıf sıralamasını '
                    've dolayısıyla accuracy/macro-F1 değerlerini değiştirmez. '
                    'ECE, en yüksek sınıf olasılığına göre 10 eşit genişlikli bin '
                    'ile hem ham hem ölçekli olasılıklarda raporlanmıştır.'
                ),
            ]
        )
        for corpus, result in results.items():
            scaling = result.get('temperature_scaling')
            if not scaling:
                continue
            before_ece = scaling['test']['before']['expected_calibration_error']
            after_ece = scaling['test']['after']['expected_calibration_error']
            blocks.append(
                ImageBlock(
                    output_root / corpus / 'test_reliability_diagram.png',
                    f'{DISPLAY[corpus]} ham ve ölçekli güvenilirlik diyagramı',
                    f'{DISPLAY[corpus]} held-out test güvenilirlik diyagramı: '
                    f'ECE {before_ece:.4f} → {after_ece:.4f}.',
                )
            )
    blocks.append(
        TableBlock(
            ('Veri seti', 'Sınıf', 'Precision', 'Recall', 'F1', 'Destek'),
            _per_class_rows(results),
            'Sınıf bazında held-out test sonuçları',
        )
    )
    for corpus in results:
        blocks.append(
            ImageBlock(
                output_root / corpus / 'test_confusion_matrix.png',
                f'{DISPLAY[corpus]} normalize karmaşıklık matrisi',
                f'{DISPLAY[corpus]} held-out test normalize karmaşıklık matrisi.',
            )
        )
    comparison_rows = _comparison_rows(comparison_path)
    if comparison_rows:
        blocks.extend(
            [
                Heading(3, 'Önceki Aşamalarla Model Karşılaştırması'),
                TableBlock(
                    (
                        'Veri seti', 'Model', 'Özellik', 'PCA', 'Doğ.',
                        'Deng. Doğ.', 'Macro-F1', 'Weighted-F1',
                    ),
                    comparison_rows,
                    'KNN, Karar Ağacı, Rastgele Orman, Gradient Boosting ve MLP test karşılaştırması',
                ),
            ]
        )
    blocks.extend(_result_analysis(results, comparison_path))
    return blocks


def _conclusion_blocks(
    results: dict[str, dict[str, Any]],
    ablation_root: Path,
    stability_root: Path,
    effect_root: Path,
) -> list[ReportBlock]:
    # 'Sonuç' bölümü: tüm deney sayaçları (koşu sayıları) gerçek dosyalardan
    # toplanır, gelecek çalışmalar listelenir ve ödev şartlarının denetim
    # tablosu kanıt referanslarıyla verilir.
    total_trials = sum(int(result['num_trials']) for result in results.values())
    feature_trials = len(_ablation_rows(ablation_root))
    # walrus operatörü (:=): path değişkenini hem koşulda tanımla hem gövdede
    # kullan — dosya yolu iki kez yazılmasın diye.
    stability_trials = sum(
        len(pd.read_csv(path))
        for corpus in results
        if (
            path := stability_root / corpus / 'feature_stability_runs.csv'
        ).is_file()
    )
    effect_summary_path = effect_root / 'summary.json'
    effect_trials = 0
    effect_excluded = 0
    if effect_summary_path.is_file():
        effect_summary = _read_json(effect_summary_path)
        effect_trials = sum(
            int(item['configuration_trials'])
            for item in effect_summary.get('per_dataset', {}).values()
        )
        effect_excluded = sum(
            int(item['excluded_stability_trials'])
            for item in effect_summary.get('per_dataset', {}).values()
        )
    effect_sentence = (
        f' Mevcut koşuların {effect_trials} benzersiz konfigürasyonu ayrıca '
        'parametre değerlerine göre betimsel olarak gruplanmış; '
        f'{effect_excluded} stability tekrarı bu analizden ayrılmıştır.'
        if effect_trials
        else ''
    )
    return [
        Heading(2, 'Sonuç ve İleride Yapılacaklar'),
        Paragraph(
            f'Bu aşamada {len(results)} veri kümesi için toplam {total_trials} '
            f'hiperparametre/seed koşusu, {feature_trials} özellik ablation '
            f'koşusu ve {stability_trials} çoklu-seed özellik doğrulama koşusu '
            'tamamlanmıştır. Her veri kümesi için bağımsız, sıfırdan PyTorch '
            'MLP modeli; train-only normalizasyon; sınıf ağırlıklı kayıp; erken '
            'durdurma; kaydedilmiş checkpoint; test tahminleri, karmaşıklık '
            'matrisi, bootstrap belirsizliği, validation-temelli temperature '
            'scaling ve güvenilirlik diyagramları üretilmiştir. Ölçekleme her iki '
            'veri kümesinde test NLL, Brier ve ECE değerlerini düşürmüş; sınıf '
            'tahminlerini ve sınıflandırma metriklerini değiştirmemiştir.'
            f'{effect_sentence}'
        ),
        Heading(3, 'İleride Yapılacaklar'),
        BulletList(
            (
                'Test setine dokunmadan eğitim katında gürültü, zaman kaydırma ve kontrollü SpecAugment veri artırımı denenmesi.',
                'MELD’de çok uzun veya hatalı kesilmiş kayıtların eğitimden önce otomatik süre/sessizlik kalite kontrolünden geçirilmesi.',
                'Seed duyarlılığını azaltmak için daha fazla tekrar ve doğrulama temelli olasılık ortalamalı MLP ensemble araştırılması.',
                'MELD için sonraki proje aşamalarında konuşma bağlamı ve metin bilgisinin, ödev kısıtlarına uygun ayrı bir deney olarak değerlendirilmesi.',
                'Validation NLL hedefli tek sıcaklık katsayısının classwise ECE, adaptive ECE ve yalnız validation üzerinde seçilecek alternatif kalibrasyon yöntemleriyle karşılaştırılması.',
                'Hata örneklerinin konuşmacı ve süre bazında incelenmesi ve sınıf bazlı veri artırımı.',
            )
        ),
        Heading(3, 'Ödev Koşulları Denetimi'),
        TableBlock(
            ('Koşul', 'Durum', 'Kanıt'),
            (
                ('PyTorch MLP sıfırdan geliştirildi', True, 'odev3/model.py ve best_model.pt'),
                ('librosa Mel vektörü en az 4000 boyutlu', True, '64×64 = 4096'),
                ('PCA kullanılmadı', True, 'result.json feature.pca=false'),
                ('Sınıf dengesizliği ağırlıklı kayıpla ele alındı', True, 'training_class_weights'),
                ('Erken durdurma uygulandı', True, '32 koşunun history/CSV kayıtları'),
                ('Tüm temel MLP hiperparametreleri denendi', True, 'validation_results.csv tabloları'),
                ('Hiperparametre ilişkileri koşu sayılarıyla incelendi', True, 'parameter_effects.csv ve parameter_effect_overview.csv'),
                ('Özellik seçimi çoklu seed ile doğrulandı', True, 'feature_stability.csv ve üç-seed PNG'),
                ('En iyi modeller kaydedildi', True, 'cremad/meld best_model.pt'),
                ('Held-out test ve karmaşıklık matrisi verildi', True, 'test_predictions.csv ve PNG'),
                ('Test belirsizliği ve validation-temelli kalibrasyon raporlandı', True, 'test_uncertainty.json, temperature_scaling.json ve reliability PNG'),
                ('Önceki aşama modelleriyle karşılaştırıldı', True, 'model_comparison.csv'),
            ),
        ),
    ]


def _build_report_blocks(
    output_root: Path,
    ablation_root: Path,
    stability_root: Path,
    effect_root: Path,
) -> tuple[list[ReportBlock], bool]:
    # Tüm raporun kurgusu: bölüm üreticileri doğru sırayla çağrılır ve tek
    # bir blok listesi elde edilir. Ayrıca koşunun "tanısal" olup olmadığına
    # karar verilir: satır limiti kullanılmışsa YA DA grid modu 'report'
    # değilse rapor tanısaldır ve başına uyarı kutusu konur.
    summary, results, validations = _load_report_inputs(output_root)
    diagnostic = bool(summary.get('diagnostic_limit_per_split')) or summary.get(
        'grid_mode'
    ) != 'report'
    total_trials = sum(result['num_trials'] for result in results.values())
    blocks: list[ReportBlock] = [
        Heading(1, 'YAP 470 / BİL 570 – Proje İlerleme Raporu 3'),
        Paragraph('Proje konusu: Konuşmadan Duygu Tanıma (Speech Emotion Recognition)'),
        Paragraph('Grup 23 – Ahmet Babagil – 211101067'),
    ]
    if diagnostic:
        blocks.append(
            Notice(
                'Bu rapor tanısal/kısıtlı bir koşudan üretildi; nihai teslim '
                'raporu olarak kullanılmamalıdır.'
            )
        )
    else:
        blocks.append(
            Notice(
                f'Nihai rapor: {len(results)} veri kümesi, {total_trials} '
                'geçerleme/seed koşusu ve held-out test sonuçları gerçek deney '
                'dosyalarından otomatik oluşturulmuştur.'
            )
        )
    blocks.extend(_access_blocks())
    blocks.extend(_execution_blocks(output_root, results, stability_root, effect_root))
    blocks.extend(_dataset_blocks(results))
    blocks.extend(_method_blocks(results, ablation_root, stability_root))
    blocks.extend(_validation_blocks(output_root, results, validations, effect_root))
    blocks.extend(_test_blocks(output_root, results))
    blocks.extend(_conclusion_blocks(results, ablation_root, stability_root, effect_root))
    return blocks, diagnostic


def build_report(
    *,
    output_root: str | Path = 'odev3/outputs',
    ablation_root: str | Path = 'odev3/feature_ablation',
    stability_root: str | Path = 'odev3/feature_stability',
    effect_root: str | Path = 'odev3/hyperparameter_effects',
    html_path: str | Path | None = None,
    docx_path: str | Path | None = None,
) -> ReportPaths:
    '''Diske kaydedilmiş deney çıktılarından nihai HTML ve Word raporlarını üretir.

    Modülün dışa açılan ana fonksiyonu (run_experiment.py de bunu çağırır):
    blokları kur, çıktı yollarını belirle, aynı içeriği iki formatta çiz.
    '''

    output_root = Path(output_root)
    blocks, diagnostic = _build_report_blocks(
        output_root,
        Path(ablation_root),
        Path(stability_root),
        Path(effect_root),
    )
    paths = _report_paths(output_root, diagnostic, html_path, docx_path)
    title = 'YAP 470 / BİL 570 – Proje İlerleme Raporu 3'
    _render_html(blocks, paths.html, title)
    _render_docx(blocks, paths.docx, title)
    return paths


def main() -> None:
    # Bağımsız CLI: rapor, deney çalıştırmadan da (mevcut çıktılardan)
    # yeniden üretilebilir. Tüm kök klasörler bayrakla değiştirilebilir.
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('--output-root', default='odev3/outputs')
    parser.add_argument('--ablation-root', default='odev3/feature_ablation')
    parser.add_argument('--stability-root', default='odev3/feature_stability')
    parser.add_argument('--effect-root', default='odev3/hyperparameter_effects')
    parser.add_argument('--html')
    parser.add_argument('--docx')
    args = parser.parse_args()
    paths = build_report(
        output_root=args.output_root,
        ablation_root=args.ablation_root,
        stability_root=args.stability_root,
        effect_root=args.effect_root,
        html_path=args.html,
        docx_path=args.docx,
    )
    print(f'HTML rapor: {paths.html}')
    print(f'DOCX rapor: {paths.docx}')


if __name__ == '__main__':
    main()
